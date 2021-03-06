
import os
import shutil

from docx.shared import Inches
from docx import Document

from PyQt5.QtWidgets import QFileDialog, QLabel, QProgressBar, QTextEdit, QVBoxLayout, QWidget, QCheckBox, QPushButton
from PyQt5.QtGui import QPixmap, QMovie, QFont
from PyQt5.QtCore import QEventLoop, QThread, QTimer, pyqtSignal, QThreadPool

from project.TaggableImageWidget import TaggableImageWidget
from project.TagWidget import TagWidget
from project.Worker import Worker

def put_tags_on_image_widget(image_widget: TaggableImageWidget, tags):
    for idx, tag in enumerate(tags, start=1):
        tag_widget = TagWidget(str(idx), "", tag["x"], tag["y"])
        tag_widget.adjust_to_size(image_widget.size(), 0.015/2)

        image_widget.add_widget(tag_widget)

def apply_settings_to_tags(tags, settings):
    results = []
    SEPARATORS_TO_NORMALIZE = [";", ":", "/", " ,", " - ", " -", "- "]
    for tag in tags:
        if settings["delete_dots"]:
            tag["text"] = tag["text"].replace(".", " ")
        if settings["normalize_separators"]:
            for separator in SEPARATORS_TO_NORMALIZE:
                tag["text"] = tag["text"].replace(separator, ", ")
        if settings["normalize_whitespaces"]:
            tag["text"] = " ".join(tag["text"].split())
        if settings["normalize_letters_to_lowercase"]:
            tag["text"] = tag["text"].lower()
        results.append(tag)
    return results
    

def generate_doc(workspace_path, data_dictionary, settings, progress_callback):
    document = Document()

    temp_dir_path = os.path.join(workspace_path, "_trener_dane_temp_")
    if os.path.exists(temp_dir_path):
        shutil.rmtree(temp_dir_path, ignore_errors=True)
        os.mkdir(temp_dir_path)
    else:
        os.mkdir(temp_dir_path)

    for imagename, tags in data_dictionary.items():
        if (len(tags) == 0):
            continue

        tags = apply_settings_to_tags(tags, settings)

        progress_callback.emit(imagename)

        pixmap_src = os.path.join(workspace_path, imagename)
        if not os.path.exists(pixmap_src):
            print("{0} has tagged structures but the image file does not exist".format(imagename))
            continue

        pixmap = QPixmap(pixmap_src)

        image_widget = TaggableImageWidget(pixmap)
        image_widget.setMinimumSize(pixmap.size())

        import pprint
        pprint.pprint(tags)

        put_tags_on_image_widget(image_widget, tags)

        image_path = os.path.join(temp_dir_path, imagename)

        image_widget.save(image_path)
        document.add_picture(image_path, Inches(7.5))
        document.add_paragraph(imagename)
        table = document.add_table(len(tags), cols=2)
        table.style = 'TableGrid'
        table.allow_autofit = False
        table.columns[0].width = Inches(0.5)
        table.columns[1].width = Inches(7)
        for idx, tag in enumerate(tags):
            table.rows[idx].cells[0].text = str(idx+1)
            table.rows[idx].cells[0].width = Inches(0.5)
            table.rows[idx].cells[1].text = tag["text"]
            table.rows[idx].cells[1].width = Inches(7)

        document.add_page_break()

    sections = document.sections
    for section in sections:
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    return document

def checked(checkbox):
    checkbox.setChecked(True)
    return checkbox

class GenerateDocSettingsWidget(QWidget):
    def __init__(self, app):
        super().__init__()
        self.app = app
        self.generate_doc_widget: GenerateDocWidget(app) = None
        self.init_ui()

    def init_ui(self):
        self.setWindowTitle("Ustawienia generowania dokumentu")
        self.layout = QVBoxLayout()

        self.delete_dots = checked(QCheckBox("Usunąć kropki?"))
        self.layout.addWidget(self.delete_dots)

        self.normalize_separators = checked(QCheckBox("Zamienić wszystkie dziwne znaki na przecinki? ('/' -> ',' etc)"))
        self.layout.addWidget(self.normalize_separators)

        self.normalize_whitespaces = checked(QCheckBox("Ujednolicić białe znaki? ('bla,bla' -> 'bla, bla' etc)"))
        self.layout.addWidget(self.normalize_whitespaces)

        self.normalize_letters_to_lowercase = checked(QCheckBox("Zamienić wszystkie znaki na małe? ('Ulna DEX' -> 'ulna dex' etc)"))
        self.layout.addWidget(self.normalize_letters_to_lowercase)

        self.generate_btn = QPushButton("Wygeneruj")
        self.generate_btn.clicked.connect(self.start)
        self.layout.addWidget(self.generate_btn)
        
        self.setLayout(self.layout)
    
    def start(self):
        self.generate_doc_widget = GenerateDocWidget(self.app)
        self.generate_doc_widget.show()
        self.generate_doc_widget.start({
            "delete_dots": self.delete_dots.isChecked(),
            "normalize_separators": self.normalize_separators.isChecked(),
            "normalize_whitespaces": self.normalize_whitespaces.isChecked(),
            "normalize_letters_to_lowercase": self.normalize_letters_to_lowercase.isChecked()
        })
        self.close()

class GenerateDocWidget(QWidget):
    def __init__(self, app):
        super().__init__()
        self.app = app
        self.threadpool = QThreadPool()
        self.images = None
        self.init_ui()

    def init_ui(self):
        self.setWindowTitle("Generowanie dokumentu w toku...")

        self.layout = QVBoxLayout()

        self.label = QLabel()
        self.movie = QMovie(os.path.join("resources", "waiting.gif"))
        self.movie.start()
        self.label.setMovie(self.movie)
        self.layout.addWidget(self.label)

        self.progressbar = QProgressBar()
        self.layout.addWidget(self.progressbar)

        self.textedit = QTextEdit()
        self.textedit.setEnabled(False)
        self.layout.addWidget(self.textedit)
        self.setLayout(self.layout)

    def start(self, settings):
        self.images = list(self.app.workspace_structures.keys())
        worker = Worker(self.generate_doc, settings=settings)
        worker.signals.result.connect(self.on_document_ready)
        worker.signals.progress.connect(self.on_progress)
        self.threadpool.start(worker)

    def generate_doc(self, progress_callback, settings):
        print(settings)
        return generate_doc(self.app.workspace_path, self.app.workspace_structures, settings, progress_callback)

    def on_progress(self, imagename):
        self.textedit.append("Przygotowuję {}...".format(imagename))
        self.progressbar.setValue(100.0 * float(self.images.index(imagename)+1) / float(len(self.images)))
        print("on progress:", imagename)

    def on_document_ready(self, document):
        self.hide()
        document_path = QFileDialog.getSaveFileName(None, "Zapisz dokument", "oznaczone struktury.docx", "docx (*.docx)")
        if document_path[0] != "":
            document.save(document_path[0])
        self.close()

