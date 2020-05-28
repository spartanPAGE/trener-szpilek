import os

from PyQt5.QtWidgets import QFileDialog, QMessageBox

from ast import literal_eval

from PyQt5.QtGui import QPixmap

from docx import Document

from project.TaggableImageWidget import TaggableImageWidget
from project.TagWidget import TagWidget
from docx.shared import Inches


def put_tags_on_image_widget(image_widget: TaggableImageWidget, tags):
    for idx, tag in enumerate(tags, start=1):
        tag_widget = TagWidget(str(idx), "", tag["x"], tag["y"])
        tag_widget.adjust_to_size(image_widget.size(), 0.015)

        image_widget.add_widget(tag_widget)

def generate_doc(workspace_path, data_dictionary):
    document = Document()

    for imagename, tags in data_dictionary.items():
        if (len(tags) == 0):
            continue

        pixmap = QPixmap(os.path.join(workspace_path, imagename))

        image_widget = TaggableImageWidget(pixmap)
        image_widget.setMinimumSize(pixmap.size())
        put_tags_on_image_widget(image_widget, tags)

        temp_dir_path = os.path.join(workspace_path, "_dane_trenera_temp_")
        if not os.path.exists(temp_dir_path):
            os.mkdir(temp_dir_path)

        image_path = os.path.join(temp_dir_path, imagename)

        image_widget.save(image_path)
        document.add_picture(image_path, Inches(7.5))
        table = document.add_table(len(tags), cols=2)
        table.style = 'TableGrid'
        table.allow_autofit = False
        table.columns[0].width = Inches(0.5)
        table.columns[1].width = Inches(7)
        for idx, tag in enumerate(tags):
            table.rows[idx].cells[0].text = str(idx+1)
            table.rows[idx].cells[0].width = Inches(0.5)
            table.rows[idx].cells[1].text = tag["text"]
            table.rows[idx].cells[1].width = Inches(7)
            print(str(idx+1), tag["text"])

        document.add_page_break()

    sections = document.sections
    for section in sections:
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    document_path = QFileDialog.getSaveFileName(None, "Zapisz dokument", "oznaczone struktury.docx", "docx (*.docx)")
    if document_path[0] != "":
        document.save(document_path[0])

def save_structures(workspace_path, data_dictionary, data_directory="_trener_dane_"):
    dir_path = os.path.join(workspace_path, data_directory)
    if not os.path.isdir(dir_path):
        os.mkdir(dir_path)

    for imagename, data in data_dictionary.items():
        with open(os.path.join(dir_path, imagename + ".txt"), "w", encoding="utf-8") as f:
            f.write('{}'.format(data))

def load_structures(workspace_path="", data_directory="_trener_dane_"):
    def show_error_messagebox():
        msg = QMessageBox()
        msg.setIcon(QMessageBox.Critical)
        msg.setText("Error")
        msg.setInformativeText('Dane trenera zakodowane są w nieznany sposób. Spróbuj zamienić je na UTF-8 i spróbuj ponownie')
        msg.setWindowTitle("Błąd odczytu")
        msg.exec_()

    data_dictionary = {}

    if workspace_path:
        dir_path = os.path.join(workspace_path, data_directory)
    else:
        dir_path = data_directory

    if os.path.exists(dir_path) and os.path.isdir(dir_path):
        for path in os.listdir(dir_path):
            try:
                with open(os.path.join(dir_path, path), 'r', encoding='utf-8') as f:
                    data_dictionary[path.replace(".txt", "")] = literal_eval(f.read())
            except UnicodeDecodeError:
                try:
                    with open(os.path.join(dir_path, path), 'r', encoding='windows-1250') as f:
                        data_dictionary[path.replace(".txt", "")] = literal_eval(f.read())
                except UnicodeDecodeError:
                    show_error_messagebox()
                    break
    return data_dictionary
